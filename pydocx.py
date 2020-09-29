#This section of code imports all of the external libaries this module needs
from docx import Document
import os

#This subroutine is used to create a new document if a document with the selected filename does not already exist
def CreateOrEdit(filename, route):
    folder = "Travel Plans"
    exists = CheckForExistingFiles(filename, folder)
    if exists != True:
        CreateNewDoc(filename, folder, route)

    return exists

#This subroutine checks the folder "Travel plans" for any files with the same file name as the one the user entered
def CheckForExistingFiles(filename, folder):
    filename = filename + ".docx"
    path = os.getcwd() + "\\" + folder
    folderContents = os.listdir(path)
    fileExists = False
    for file in folderContents:
        if filename == file:
            fileExists = True
            break

    return fileExists

#This subroutine creates a new document and then adds the users instructions to the word document
def CreateNewDoc(filename, folder, route):
    d = Document()
    title = filename
    filename = filename + ".docx"
    d.add_heading(title, level = 0)
    d.add_heading(route[0][0] + " to " + route[-1][1], level = 1)
    for i in route:
        p = d.add_paragraph("{0} to {1} on {2} line for {3} minutes".format(i[0], i[1], i[2], i[3]))

    d.save(folder + "/" + filename)

#This subroutine opens an existing document and then appends the routes instructions to the bottom of the document
def EditExistingDoc(filename, folder, route):
    filename = filename + ".docx"
    d = Document(folder + "/" + filename)
    d.add_paragraph("")
    d.add_heading(route[0][0] + " to " + route[-1][1], level = 1)
    for i in route:
        p = d.add_paragraph("{0} to {1} on {2} line for {3} minutes".format(i[0], i[1], i[2], i[3]))

    d.save(folder + "/" + filename)
